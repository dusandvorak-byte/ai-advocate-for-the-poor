from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path('/workspace/scratch/491cfa754b55')
UPLOAD = ROOT / 'upload'
OUT = ROOT / 'deliverables'

SOURCE = {
    'ncoz': UPLOAD / 'Městskému soudu dne 18.července 2026 - NCOZ -extrémní naléhavost.docx',
    'msp': UPLOAD / 'Městskému soudu dne 18.července 2026 - MSP - Stížnost ministru spravedlnosti - naléhavost.docx',
    'president': UPLOAD / 'Prezidentu republiky dne 18. července 2026 - nové důkazy a doplnění žádosti - naléhavost - stížnost.docx',
}

OUTPUT = {
    'ncoz': OUT / 'NCOZ_doplneno_KSZ_2026-07-19.docx',
    'msp': OUT / 'MSp_doplneno_KSZ_2026-07-19.docx',
    'president': OUT / 'Prezident_doplneno_KSZ_2026-07-19.docx',
}

SOURCE_DESCRIPTION = (
    'sdělením Krajského státního zastupitelství v Ostravě ze dne '
    '8. července 2026, sp. zn. 4 KZN 7116/2026-45, doručeným dne 16. července 2026'
)

SOURCE_QUOTE = (
    '„k vyhodnocení a přijetí odpovídajícího opatření“'
)


def find_one(doc, startswith):
    matches = [p for p in doc.paragraphs if p.text.strip().startswith(startswith)]
    if len(matches) != 1:
        raise ValueError(f'Expected one paragraph starting with {startswith!r}, found {len(matches)}')
    return matches[0]


def set_text_preserve_format(paragraph, text):
    if not paragraph.runs:
        paragraph.add_run(text)
        return paragraph
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ''
    return paragraph


def clone_before(target, template, text):
    cloned = deepcopy(template._p)
    target._p.addprevious(cloned)
    paragraph = Paragraph(cloned, target._parent)
    return set_text_preserve_format(paragraph, text)


def append_clone(doc, template, text):
    cloned = deepcopy(template._p)
    doc._body._body.append(cloned)
    paragraph = doc.paragraphs[-1]
    return set_text_preserve_format(paragraph, text)


def set_image_alt_texts(doc, descriptions):
    image_properties = doc._element.xpath('.//wp:docPr')
    if len(image_properties) != len(descriptions):
        raise ValueError(
            f'Expected {len(descriptions)} images for alt text, found {len(image_properties)}'
        )
    for image_property, description in zip(image_properties, descriptions):
        image_property.set('descr', description)


def build_ncoz():
    doc = Document(SOURCE['ncoz'])

    current_state_target = find_one(doc, 'MSZ v Praze dne 17. července 2026')
    clone_before(
        current_state_target,
        current_state_target,
        'KSZ v Ostravě sdělením ze dne 8. července 2026, doručeným dne 16. července 2026, '
        'oznámilo postoupení části podání týkající se postupu OKTE Frýdek-Místek na OSZ '
        'Frýdek-Místek; nová spisová značka ani věcný výsledek dosud nejsou doloženy;'
    )

    prague_heading = find_one(doc, '2. Pražská větev')
    long_heading = find_one(doc, '3. Dlouhodobá institucionální vědomost')
    set_text_preserve_format(prague_heading, '3. Pražská větev')
    set_text_preserve_format(long_heading, '4. Dlouhodobá institucionální vědomost')
    silesian_heading = clone_before(prague_heading, prague_heading, '2. Slezská větev')
    # Keep the new heading with its explanation.  In the source layout it would
    # otherwise be orphaned at the foot of a page, separated from the evidence.
    silesian_heading.paragraph_format.page_break_before = True
    silesian_heading.paragraph_format.keep_with_next = True
    clone_before(
        prague_heading,
        find_one(doc, 'VSZ v Olomouci proto dne 16. července 2026'),
        'Krajské státní zastupitelství v Ostravě sdělením ze dne 8. července 2026, '
        'sp. zn. 4 KZN 7116/2026-45, doručeným dne 16. července 2026, '
        'oznámilo, že část podání ze dne 27. května 2026 směřující vůči postupu '
        f'OKTE Frýdek-Místek postoupilo OSZ Frýdek-Místek {SOURCE_QUOTE}.'
    )
    clone_before(
        prague_heading,
        find_one(doc, 'VSZ v Olomouci proto dne 16. července 2026'),
        'Listina dokládá pouze vznik a pokračování této procesní větve. Nedokládá '
        'potvrzení nesprávnosti či nezákonnosti postupu OKTE ani přijetí konkrétního opatření. '
        'Dosud není doložena nová spisová značka OSZ ani výsledek vyhodnocení. Odkaz na přezkum '
        'cestou státního zástupce proto zatím nepředstavuje doloženou dokončenou nápravu.'
    )

    evidence_target = find_one(doc, 'sdělením MSZ v Praze ze dne 17. července 2026')
    clone_before(
        evidence_target,
        evidence_target,
        'sdělením Krajského státního zastupitelství v Ostravě ze dne 8. července 2026, '
        'sp. zn. 4 KZN 7116/2026-45, doručeným dne 16. července 2026;'
    )

    urgency_target = find_one(doc, 'biologickou úrodu ani výzkumnou sezónu')
    clone_before(
        urgency_target,
        urgency_target,
        'nově doložená slezská větev zůstává bez známé spisové značky OSZ a bez věcného výsledku;'
    )

    attachment_template = find_one(doc, 'Podnět ÚOOÚ dne 17.července 2026')
    append_clone(
        doc,
        attachment_template,
        'Sdělení Krajského státního zastupitelství v Ostravě ze dne 8. července 2026, '
        'sp. zn. 4 KZN 7116/2026-45, doručené dne 16. července 2026.'
    )

    doc.save(OUTPUT['ncoz'])


def build_msp():
    doc = Document(SOURCE['msp'])

    section_xv = find_one(doc, 'XV. Jiná řízení nezbavují')
    clone_before(
        section_xv,
        find_one(doc, 'Dánský podklad je významný'),
        'Dne 16. července 2026 bylo žalobci doručeno sdělení Krajského státního zastupitelství '
        'v Ostravě ze dne 8. července 2026, sp. zn. 4 KZN 7116/2026-45, podle něhož byla část '
        'podání týkající se postupu OKTE Frýdek-Místek postoupena Okresnímu státnímu zastupitelství '
        f've Frýdku-Místku {SOURCE_QUOTE}. Jde o novou procesní skutečnost dokládající otevřenou '
        'forenzní větev. Listina nepotvrzuje pochybení OKTE ani výsledek dohledu.'
    )

    responsibility_target = find_one(doc, 'Existence jiných řízení proto není důvodem')
    clone_before(
        responsibility_target,
        responsibility_target,
        'Nové postoupení nepřenáší odpovědnost KSZ nebo OSZ na Ministerstvo spravedlnosti. '
        'Žalobce je předkládá pouze k doložení aktuálního stavu a naléhavosti koordinovaného '
        'vypořádání, aniž mění žalovaného nebo petit.'
    )

    urgency_target = find_one(doc, 'žalovaný dosud nesdělil, zda zahájil')
    clone_before(
        urgency_target,
        urgency_target,
        'nová procesní větev týkající se OKTE Frýdek-Místek zůstává bez doloženého '
        'věcného výsledku;'
    )

    evidence_target = find_one(doc, 'odpověď Kriminalistického ústavu ze dne 25. června 2026')
    clone_before(
        evidence_target,
        evidence_target,
        'sdělení Krajského státního zastupitelství v Ostravě ze dne 8. července 2026, '
        'sp. zn. 4 KZN 7116/2026-45, doručené dne 16. července 2026;'
    )

    attachment_intro = find_one(doc, 'Přílohy tučně vyznačené vztahující se primárně')
    set_text_preserve_format(
        attachment_intro,
        'Níže uvedené přílohy se primárně vztahují k zásahové žalobě proti NCOZ, '
        'sp. zn. 18 A 17/2026, a jsou současně určeny k předložení do tohoto spisu. Žalobce '
        'žádá, aby je soud po jejich prokazatelném doručení hodnotil v obou souvisejících věcech '
        've vzájemné souvislosti. Toto podání netvrdí, že byly soudu doručeny dříve, než to '
        'bude doloženo doručenkou.'
    )
    attachment_target = find_one(doc, '13. Podnět ÚOOÚ dne 17. července 2026')
    clone_before(
        attachment_target,
        attachment_target,
        '13. Sdělení Krajského státního zastupitelství v Ostravě ze dne 8. července 2026, '
        'sp. zn. 4 KZN 7116/2026-45, doručené dne 16. července 2026.'
    )
    set_text_preserve_format(attachment_target, attachment_target.text.replace('13.', '14.', 1))

    set_image_alt_texts(
        doc,
        ['Fotografie Dušana Dvořáka sedícího mezi vzrostlými rostlinami konopí.']
    )
    doc.save(OUTPUT['msp'])


def build_president():
    doc = Document(SOURCE['president'])

    paragraph_34 = find_one(doc, 'Dne 18. července 2026 byla oběma řízením')
    set_text_preserve_format(
        paragraph_34,
        'Dne 18. července 2026 byla připravena dvě další mimořádně naléhavá doplnění '
        'určená pro řízení senátu 18 A Městského soudu v Praze. Ke dni sepsání tohoto '
        'podání jejich doručení soudu není tvrzeno a musí být doloženo samostatnými doručenkami.'
    )
    paragraph_35 = find_one(doc, 'Ve věci sp. zn. 18 A 23/2026 proti Ministerstvu spravedlnosti bylo doloženo')
    set_text_preserve_format(
        paragraph_35,
        'Připravené doplnění ve věci sp. zn. 18 A 23/2026 proti Ministerstvu spravedlnosti dokládá, '
        'že ani po stížnosti ministrovi spravedlnosti ze dne 15. července 2026 nebylo přezkoumatelně '
        'sděleno, zda ministerstvo zahájilo dohledový nebo odborný přezkum znaleckých postupů. '
        'Připravené doplnění ve věci sp. zn. 18 A 17/2026 proti NCOZ dokládá, že podání '
        'z období 25. dubna až 12. května 2026 byla uzavřena bez existujícího záznamu o individuálním '
        'a odborném hodnocení a že tento stav neodstranilo ani rozhodnutí Ministerstva vnitra ze dne '
        '16. července 2026, č. j. MV-97289-3/TP-2026. Obě žaloby mají rozdílné žalované a rozdílný '
        'předmět, avšak společně dokládají bezprostřední riziko opakování roku 2019.'
    )

    section_iii = find_one(doc, 'III. Podstata nezodpovězených otázek')
    clone_before(
        section_iii,
        paragraph_35,
        'Chronologie byla nově doplněna o sdělení Krajského státního zastupitelství v Ostravě '
        'ze dne 8. července 2026, sp. zn. 4 KZN 7116/2026-45, doručené dne 16. července 2026. '
        'Část podání vztahující se k postupu OKTE Frýdek-Místek byla postoupena Okresnímu '
        f'státnímu zastupitelství ve Frýdku-Místku {SOURCE_QUOTE}. Tato listina nepotvrzuje pochybení '
        'ani přijetí opatření; dokládá, že další související procesní větev zůstává otevřená '
        'bez doložené spisové značky OSZ a bez věcného výsledku.'
    )

    ps_heading = find_one(doc, '18. července 2026 – dvě mimořádně naléhavá doplnění')
    set_text_preserve_format(ps_heading, '18. července 2026 – dvě připravená mimořádně naléhavá doplnění zásahových žalob')
    ps_text = find_one(doc, 'Ve věci 18 A 23/2026 proti Ministerstvu spravedlnosti byla doložena')
    set_text_preserve_format(
        ps_text,
        'Připravené doplnění ve věci 18 A 23/2026 proti Ministerstvu spravedlnosti má doložit '
        'trvající absenci přezkoumatelné odpovědi, zda byl zahájen dohled nad znalci a znaleckými '
        'ústavy. Připravené doplnění ve věci 18 A 17/2026 proti NCOZ má doložit uzavření nových '
        'podání bez záznamu o jejich odborném posouzení. Obě podání žádají přednostní ochranu; '
        'jejich skutečné odeslání a doručení bude doloženo samostatnými doručenkami.'
    )

    kpr_heading = find_one(doc, '14.–16. července 2026 – KPR')
    clone_before(kpr_heading, kpr_heading, '8./16. července 2026 – KSZ Ostrava a OSZ Frýdek-Místek')
    clone_before(
        kpr_heading,
        find_one(doc, 'KPR byla seznámena s předžalobní výzvou NSZ'),
        'KSZ Ostrava sdělilo postoupení části podání týkající se postupu OKTE Frýdek-Místek '
        'na OSZ Frýdek-Místek. Postoupení není potvrzením pochybení; vytváří nový otevřený '
        'procesní uzel bez doložené spisové značky OSZ a bez věcného výsledku.'
    )

    attachment_target = find_one(doc, 'První ústavní stížnost ze dne 24. února 2012')
    clone_before(
        attachment_target,
        find_one(doc, 'Stížnost ministrovi vnitra Mgr. Lubomíru Metnarovi'),
        'Sdělení Krajského státního zastupitelství v Ostravě ze dne 8. července 2026, '
        'sp. zn. 4 KZN 7116/2026-45, doručené dne 16. července 2026.'
    )

    set_image_alt_texts(
        doc,
        [
            'Plakát festivalu Noc básníků v Ospělově dne 21. srpna.',
            'Stránka dětské encyklopedie s obrázkem a popisem konopí.'
        ]
    )
    doc.save(OUTPUT['president'])


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    build_ncoz()
    build_msp()
    build_president()
    for path in OUTPUT.values():
        print(path)


if __name__ == '__main__':
    main()
