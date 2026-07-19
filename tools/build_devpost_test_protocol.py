from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path('/workspace/scratch/491cfa754b55')
OUT = ROOT / 'deliverables'
OUTPUT_CS = OUT / 'AI_Advocate_Testovaci_Protokol_CZ_2026-07-19.docx'
OUTPUT_EN = OUT / 'AI_Advocate_Devpost_Test_Protocol_EN_2026-07-19.docx'

NAVY = '17324D'
BLUE = '176B87'
LIGHT_BLUE = 'EAF4F7'
RED = 'A12A2A'
LIGHT_RED = 'FBECEC'
AMBER = 'A65F00'
LIGHT_AMBER = 'FFF4DE'
GREEN = '27623A'
LIGHT_GREEN = 'EAF5ED'
GRAY = '5B6670'
LIGHT_GRAY = 'F2F4F5'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn('w:shd'))
    if shading is None:
        shading = OxmlElement('w:shd')
        tc_pr.append(shading)
    shading.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=100, start=110, bottom=100, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for margin, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{margin}'))
        if node is None:
            node = OxmlElement(f'w:{margin}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('AI Advocate for the Poor  |  19 July 2026  |  ')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_hyperlink(paragraph, label, url):
    relationship = paragraph.part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), relationship)
    run = OxmlElement('w:r')
    props = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), BLUE)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    props.extend([color, underline])
    run.append(props)
    text = OxmlElement('w:t')
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement('w:tblHeader')
    repeat.set(qn('w:val'), 'true')
    tr_pr.append(repeat)


def keep_table_rows_together(table):
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        tr_pr.append(cant_split)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Aptos'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ('Title', 28, NAVY),
        ('Subtitle', 14, BLUE),
        ('Heading 1', 19, NAVY),
        ('Heading 2', 13, BLUE),
        ('Heading 3', 11, NAVY),
    ]:
        style = styles[name]
        style.font.name = 'Aptos Display'
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != 'Subtitle'
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if name != 'Title' else 0)
        style.paragraph_format.space_after = Pt(5)

    for doc_section in doc.sections:
        header = doc_section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.text = 'OPENAI BUILD WEEK · DEVPOST · PRE-SUBMISSION TEST'
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            run.font.size = Pt(7.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(BLUE)
        doc_section.footer.is_linked_to_previous = False
        add_page_number(doc_section.footer.paragraphs[0])


def add_cover(doc, language):
    cs = language == 'cs'
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(38)
    run = p.add_run('AI ADVOCATE FOR THE POOR')
    run.bold = True
    run.font.name = 'Aptos Display'
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PROTOKOL TESTOVÁNÍ PŘED PODÁNÍM PŘIHLÁŠKY' if cs else 'PRE-SUBMISSION TEST PROTOCOL')
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(15)
    p.add_run('OpenAI Build Week · Devpost\n').bold = True
    p.add_run('Submission 1098095 · 19 July 2026\n')
    p.add_run('Controlled internal-archive update and safe outside-input test' if not cs else 'Řízená aktualizace interního archivu a bezpečný test vnějšího vstupu')

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(5.0)
    table.columns[1].width = Cm(11.2)
    header_cells = table.add_row().cells
    header_cells[0].text = 'Identifikace' if cs else 'Identification'
    header_cells[1].text = 'Údaj' if cs else 'Details'
    for cell in header_cells:
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        for header_run in cell.paragraphs[0].runs:
            header_run.bold = True
            header_run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_table_header(table.rows[0])
    entries = [
        ('Předkladatel' if cs else 'Applicant', 'Mgr. Dušan Dvořák'),
        ('Organizace' if cs else 'Organisation', 'Cannabis is The Cure, z.s. · IČO 26670232'),
        ('Trvalé bydliště' if cs else 'Permanent address', 'Educational Cannabis Clinic, Ospělov 6, 798 55, Czech Republic'),
        ('Korespondenční adresa' if cs else 'Correspondence address', 'Cannabis is The Cure, z.s., Přichystalova 180/14, 779 00 Olomouc, Czech Republic'),
        ('Telefon' if cs else 'Telephone', '+420 774 723 261'),
        ('Kontaktní e-mail' if cs else 'Contact email', 'dusandvorak@seznam.cz'),
        ('E-mail účtu Devpost' if cs else 'Devpost account email', 'dusan.dvorak@konopijelek.cz'),
    ]
    for label, value in entries:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_BLUE)
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[0].paragraphs[0].runs[0].bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run(
        'Tento protokol dokládá funkci prototypu. Není právním podáním, nepotvrzuje odeslání žádného dokumentu a nenahrazuje kontrolu člověkem.'
        if cs else
        'This protocol records prototype behaviour. It is not a legal filing, confirms no document delivery, and does not replace human review.'
    )
    run.italic = True
    run.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_page_break()


def add_callout(doc, title, body, color='blue'):
    fills = {'blue': LIGHT_BLUE, 'red': LIGHT_RED, 'amber': LIGHT_AMBER, 'green': LIGHT_GREEN}
    colors = {'blue': BLUE, 'red': RED, 'amber': AMBER, 'green': GREEN}
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_repeat_table_header(table.rows[0])
    set_cell_shading(cell, fills[color])
    set_cell_margins(cell, 140, 170, 140, 170)
    p = cell.paragraphs[0]
    r = p.add_run(title + '\n')
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(colors[color])
    p.add_run(body)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc, items):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.55)
        p.paragraph_format.space_after = Pt(3)
        number = p.add_run(f'{index}. ')
        number.bold = True
        p.add_run(item)


def add_matrix(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    table.autofit = False
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = label
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
            if row_index % 2:
                set_cell_shading(cells[i], LIGHT_GRAY)
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Cm(width)
    keep_table_rows_together(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_links(doc, cs):
    doc.add_heading('Odkazy pro porotu' if cs else 'Judge links', level=2)
    links = [
        ('Živý prototyp' if cs else 'Live prototype', 'https://dusandvorak-byte.github.io/ai-advocate-for-the-poor/'),
        ('Veřejný repozitář' if cs else 'Public repository', 'https://github.com/dusandvorak-byte/ai-advocate-for-the-poor'),
        ('Pravidla soutěže' if cs else 'Competition rules', 'https://openai.devpost.com/rules'),
    ]
    for label, url in links:
        p = doc.add_paragraph()
        p.add_run(label + ': ').bold = True
        add_hyperlink(p, 'otevřít odkaz' if cs else 'open link', url)


def add_protocol(doc, language):
    cs = language == 'cs'
    add_cover(doc, language)

    doc.add_heading('1. Účel a princip testu' if cs else '1. Purpose and test principle', level=1)
    doc.add_paragraph(
        'Cílem je ověřit, zda jediná nová listina bezpečně aktualizuje již existující archiv, aniž z procesního kroku vyrábí důkaz pochybení. Test sleduje dopad na časovou osu, lhůty, procesní pavouk, důkazní řetězec, tři dosud nepodané návrhy a okruh dalších orgánů.'
        if cs else
        'The test asks whether one new document can safely update an existing archive without turning a procedural step into proof of wrongdoing. It traces the effect on chronology, deadlines, the procedural graph, the evidence chain, three unfiled drafts, and other potentially affected authorities.'
    )
    add_callout(
        doc,
        'Vývojové pravidlo' if cs else 'Development rule',
        'Malý detail může mít na celý systém devastující dopad. Proto se každá změna propaguje napříč archivem a každé nové tvrzení musí mít zdroj, meze a stav ověření.'
        if cs else
        'A small detail can have a devastating system-wide effect. Every change is therefore propagated across the archive, and every new claim must retain its source, limitation, and verification status.',
        'blue'
    )

    doc.add_heading('2. Vstupní listina' if cs else '2. Input document', level=1)
    add_matrix(
        doc,
        ['Pole', 'Hodnota'] if cs else ['Field', 'Value'],
        [
            ('Vydal' if cs else 'Issuer', 'Krajské státní zastupitelství v Ostravě' if cs else 'Regional Public Prosecutor’s Office in Ostrava'),
            ('Datum listiny' if cs else 'Document date', '8 July 2026'),
            ('Datum doručení v archivu' if cs else 'Receipt date recorded in archive', '16 July 2026'),
            ('Veřejná značka' if cs else 'Public reference', '4 KZN 1000/2026-10 — anonymizováno / anonymised'),
            ('Klasifikace' if cs else 'Classification', 'Procesní postoupení části podání' if cs else 'Procedural referral of part of a submission'),
        ],
        [4.6, 12.0]
    )
    p = doc.add_paragraph()
    p.add_run('Rozhodná citace: ' if cs else 'Operative quotation: ').bold = True
    p.add_run('„část Vašeho podání směřující vůči postupu policejního orgánu – OKTE Frýdek-Místek byla postoupena Okresnímu státnímu zastupitelství ve Frýdku-Místku k vyhodnocení a přijetí odpovídajícího opatření“').italic = True
    if not cs:
        p = doc.add_paragraph('English translation: “the part of your submission concerning the conduct of the police authority — OKTE Frýdek-Místek — was referred to the District Public Prosecutor’s Office in Frýdek-Místek for assessment and adoption of an appropriate measure”.')
        p.runs[0].italic = True

    doc.add_heading('3. Výsledek významu a lhůty' if cs else '3. Meaning and deadline result', level=1)
    add_matrix(
        doc,
        ['Kontrola', 'Výsledek'] if cs else ['Check', 'Result'],
        [
            ('Co dokládá' if cs else 'What it proves', 'Postoupení vymezené části podání; vznik otevřené větve OSZ Frýdek-Místek; účel vyhodnocení a případného opatření.' if cs else 'Referral of a defined part of the submission; creation of an open Frýdek-Místek District Office branch; assessment and possible action as its stated purpose.'),
            ('Co nedokládá' if cs else 'What it does not prove', 'Nezákonnost nebo nesprávnost OKTE; přijetí opatření; věcný výsledek; změnu žalovaného nebo petitu.' if cs else 'Unlawful or incorrect OKTE conduct; adoption of a measure; a merits outcome; any change of defendant or relief sought.'),
            ('Zákonná lhůta' if cs else 'Statutory deadline', 'Z textu nebyla zjištěna žádná povinná lhůta k reakci.' if cs else 'No mandatory response deadline was identified in the text.'),
            ('Kontrolní bod' if cs else 'Monitoring checkpoint', '30. července 2026 — dobrovolně, 14 kalendářních dnů od doručení; nejde o právní lhůtu.' if cs else '30 July 2026 — voluntary, 14 calendar days after receipt; not a legal deadline.'),
            ('Chybějící údaje' if cs else 'Missing data', 'Nová spisová značka OSZ, rozsah převzaté části, věcný výsledek a případné opatření.' if cs else 'New District Office reference, scope of the referred part, merits outcome, and any measure.'),
        ],
        [4.2, 12.4]
    )
    add_callout(
        doc,
        'Zachycená chyba' if cs else 'Prevented error',
        'Systém odmítl vydávat praktický čtrnáctidenní kontrolní interval za zákonnou lhůtu.'
        if cs else
        'The system refused to present a practical fourteen-day checkpoint as a statutory deadline.',
        'green'
    )

    doc.add_heading('4. Aktualizace procesního pavouka' if cs else '4. Procedural-graph update', level=1)
    add_bullets(doc, [
        'Přidán nový uzel OSZ Frýdek-Místek a vztah „postoupeno k vyhodnocení a případnému opatření“.' if cs else 'Added a new Frýdek-Místek District Public Prosecutor node and a “referred for assessment and possible action” relationship.',
        'Uzel má stav „otevřený — výsledek neznámý“, nikoli „pochybení potvrzeno“.' if cs else 'The node is marked “open — outcome unknown”, not “wrongdoing confirmed”.',
        'Události 8. a 16. července 2026 byly odděleny jako datum vydání a datum doručení.' if cs else '8 and 16 July 2026 are recorded separately as issue and receipt dates.',
        'Zachována vazba na důkazní osu OKTE, ale bez automatického sloučení územně a procesně odlišných řízení.' if cs else 'The link to the OKTE evidence axis is preserved without automatically merging territorially or procedurally distinct proceedings.',
    ])

    doc.add_heading('5. Dopad na tři nepodané návrhy' if cs else '5. Impact on three unfiled drafts', level=1)
    add_matrix(
        doc,
        ['Semafor', 'Dokument', 'Výstup'] if cs else ['Signal', 'Document', 'Output'],
        [
            ('ČERVENÁ' if cs else 'RED', 'Zásahová žaloba proti NCOZ, 18 A 17/2026', 'Doplnit novou slezskou větev, důkazní návrh, naléhavost a přílohy; výslovně uvést, že postoupení není potvrzením pochybení.' if cs else 'Add the Silesian branch, evidence proposal, urgency point, and exhibit; expressly state that referral is not a finding of wrongdoing.'),
            ('ORANŽOVÁ' if cs else 'AMBER', 'Zásahová žaloba proti Ministerstvu spravedlnosti, 18 A 23/2026', 'Stručně doplnit chronologii, otevřenou forenzní větev a mez odpovědnosti; neměnit žalovaného ani petit.' if cs else 'Briefly add chronology, the open forensic branch, and the responsibility boundary; do not change defendant or relief sought.'),
            ('ORANŽOVÁ' if cs else 'AMBER', 'Dopis prezidentu republiky, KPR 4093/2026', 'Doplnit chronologii, procesní pavouk, P. S. a přílohy; neoznačit listinu za věcný výsledek.' if cs else 'Update chronology, procedural graph, postscript, and exhibits; do not describe the document as a merits outcome.'),
        ],
        [2.4, 5.8, 8.4]
    )

    doc.add_heading('6. Kritické detaily zachycené napříč návrhy' if cs else '6. Critical cross-draft details detected', level=1)
    add_matrix(
        doc,
        ['Stav', 'Zjištění', 'Bezpečná oprava'] if cs else ['Status', 'Finding', 'Safe correction'],
        [
            ('BLOKOVAT' if cs else 'BLOCK', 'Dopis prezidentu používal „byla předložena“ a „bylo doloženo“, ačkoli návrhy nebyly podány.' if cs else 'The presidential letter said “was submitted” and “was documented” although the drafts had not been filed.', 'Změnit na „byla připravena“; odeslání a doručení tvrdit až podle doručenky.' if cs else 'Change to “was prepared”; assert dispatch or delivery only when supported by a delivery record.'),
            ('BLOKOVAT' if cs else 'BLOCK', 'Návrh proti Ministerstvu spravedlnosti tvrdil, že přílohy „již byly zaslány“ do jiného spisu.' if cs else 'The Ministry of Justice draft stated that exhibits “had already been sent” to another file.', 'Hodnotit je v obou věcech až po prokazatelném doručení; předchozí doručení netvrdit.' if cs else 'Ask for cross-file consideration only after documented delivery; do not claim prior delivery.'),
            ('OVĚŘENO SOUKROMĚ' if cs else 'VERIFIED PRIVATELY', 'Aktuálnost názvů, stabilních identifikátorů a historických změn byla zkontrolována v neveřejné pracovní vrstvě.' if cs else 'Current names, stable identifiers, and historical changes were checked in a non-public working layer.', 'Soutěžní protokol zveřejňuje pouze výsledek kontroly, nikoli konkrétní interní identifikátory nebo dřívější názvy.' if cs else 'The competition protocol exposes only the check result, not specific internal identifiers or former names.'),
        ],
        [2.4, 7.2, 7.0]
    )
    add_callout(
        doc,
        'Stav tří výstupů' if cs else 'Status of the three outputs',
        'Obsahové doplnění KSZ je připraveno a vizuálně zkontrolováno. Identita a změny názvů byly ověřeny v neveřejné pracovní vrstvě. Žádný dokument nebyl tímto testem odeslán.'
        if cs else
        'The KSZ content update is prepared and visually checked. Identity and name changes were verified in the non-public working layer. This test sent no document.',
        'green'
    )

    doc.add_heading('7. Semafor dalších adresátů' if cs else '7. Traffic light for other recipients', level=1)
    add_matrix(
        doc,
        ['Semafor', 'Adresát / řízení', 'Doporučení'] if cs else ['Signal', 'Recipient / proceeding', 'Recommendation'],
        [
            ('ČERVENÁ' if cs else 'RED', 'OSZ Frýdek-Místek; KSZ Ostrava; předžalobní výzva nejvyšší státní zástupkyni; stížnosti ministrovi vnitra', 'Doplnit přesnou událost, citaci a mez tvrzení; žádat značku, návaznost a věcný výsledek.' if cs else 'Add the precise event, quotation, and claim limitation; request the new reference, chain, and merits outcome.'),
            ('ORANŽOVÁ' if cs else 'AMBER', 'Ministryně spravedlnosti; MSZ Praha; VSZ Praha', 'Doplnit stručně nebo jen po ověření procesní návaznosti.' if cs else 'Add briefly or only after verifying the procedural link.'),
            ('ORANŽOVÁ — PODMÍNĚNĚ' if cs else 'AMBER — CONDITIONAL', 'VSZ Olomouc; KSZ Brno; dohledová větev KSZ Ostrava', 'Použít jen pokud spis řeší stejnou laboratorní metodiku nebo navazuje na stejné podání.' if cs else 'Use only if the file concerns the same laboratory method or the same submission chain.'),
            ('ORANŽOVÁ — OMEZENĚ' if cs else 'AMBER — LIMITED', 'Třetí zásahová žaloba proti Ministerstvu zdravotnictví, 8 Ad 9/2026', 'Uvést pouze jako kontext otevřené laboratorní větve při souhrnném doplnění; samostatné podání nyní není nutné.' if cs else 'Mention only as open-laboratory context in a consolidated supplement; no standalone filing is currently needed.'),
            ('ZELENÁ' if cs else 'GREEN', 'ÚOOÚ; RRTV; Rada ČT; generální ředitel ČT; spor s ČT; obnova trestních řízení', 'Pouze evidovat. Procesní postoupení samo není novým věcným důkazem pro tyto větve.' if cs else 'Record only. The referral is not substantive new evidence for these branches.'),
        ],
        [3.1, 6.4, 7.1]
    )
    p = doc.add_paragraph()
    p.add_run('Význam barev: ' if cs else 'Colour meaning: ').bold = True
    p.add_run('červená znamená přímou relevanci nebo blokující kontrolu, nikoli vinu či předpověď výsledku; oranžová omezenou nebo podmíněnou relevanci; zelená znamená nyní pouze evidovat.' if cs else 'red means direct relevance or a blocking check, not guilt or predicted outcome; amber means limited or conditional relevance; green means record only for now.')

    doc.add_heading('8. Druhý test: bezpečný vstup zvenčí' if cs else '8. Second test: safe outside input', level=1)
    add_numbered(doc, [
        'Porotce vloží přesně podporované soudní PDF nebo načte jeho veřejnou anonymizovanou kopii; zpracování proběhne pouze v prohlížeči a totožnost verze se ověří digitálním otiskem.' if cs else 'A judge selects the exactly supported court PDF or loads its public anonymized copy; processing occurs only in the browser and the exact version is verified by a digital fingerprint.',
        'Systém zobrazí anonymizovanou analýzu změny v dozorčí radě, oddělí obsah usnesení od tvrzení autora a zveřejní pouze odvozený důkazní výtah.' if cs else 'The system displays an anonymized analysis of the supervisory-board change, separates the order from the archive-owner statement, and publishes only a derived evidentiary extract.',
        'Výsledek ukáže stav paměti před vložením a po vložení listiny a vytvoří konkrétní řešení: časově omezený zápis, chybějící podklady a větve, do nichž se změna nemá přenést.' if cs else 'The result shows memory before and after the document was added and generates a concrete solution: a time-bounded record, the missing evidence, and branches into which the change must not be propagated.',
        'U neznámého PDF systém připravenou analýzu odmítne. U jiného textu vytvoří pouze technickou mapu metadat bez převzetí připravených právních závěrů.' if cs else 'For an unknown PDF, the system rejects the prepared analysis. Other text receives only a technical metadata map without inheriting prepared legal conclusions.',
        'Úplný originál se neodesílá ani nezveřejňuje. Veřejný výstup neobsahuje data narození, rodná čísla, adresy, podpisy ani jiné nepotřebné údaje třetích osob.' if cs else 'The complete original is neither uploaded nor published. The public output contains no dates of birth, national identification numbers, addresses, signatures, or other unnecessary third-party data.',
    ])
    add_callout(
        doc,
        'Současné omezení' if cs else 'Current limitation',
        'Vnější PDF test přesně podporuje pouze jednu verzovanou soudní listinu a její anonymizovanou odvozenou kopii. Obecné PDF/DOCX, OCR, univerzální právní analýza a automatické podávání jsou další vývojové linie.'
        if cs else
        'The outside-PDF test exactly supports only one versioned court document and its anonymized derivative. General PDF/DOCX, OCR, universal legal analysis, and automated filing are future development lines.',
        'amber'
    )

    doc.add_heading('9. Ověření kvality a přijímací kritéria' if cs else '9. Quality verification and acceptance criteria', level=1)
    add_matrix(
        doc,
        ['Kontrola', 'Výsledek'] if cs else ['Control', 'Result'],
        [
            ('Automatické testy' if cs else 'Automated tests', '54/54 úspěšných' if cs else '54/54 passed'),
            ('Zdrojová kázeň' if cs else 'Source discipline', 'Každé připravené tvrzení odkazuje na kontrolní listinu; postoupení se nemění na potvrzení pochybení.' if cs else 'Every prepared claim remains tied to the control document; referral never becomes confirmation of wrongdoing.'),
            ('Česká terminologie' if cs else 'Czech terminology', 'Pouze „státní zástupce / státní zastupitelství“; terminologická kontrola prošla.' if cs else 'Uses “státní zástupce / státní zastupitelství” in Czech; terminology control passed.'),
            ('Dokumenty Word' if cs else 'Word documents', 'Tři revidované návrhy vykresleny a vizuálně zkontrolovány; audit přístupnosti 0 závažných, 0 středních, 0 nízkých nálezů.' if cs else 'Three revised drafts rendered and visually reviewed; accessibility audit reported 0 high, 0 medium, and 0 low findings.'),
            ('Časová identita subjektů' if cs else 'Temporal entity identity', 'Aktuální názvy a stabilní identifikátory byly sledovány odděleně od historických aliasů; konkrétní údaje zůstaly neveřejné.' if cs else 'Current names and stable identifiers were tracked separately from historical aliases; specific details remained non-public.'),
            ('Anonymizace vnějšího PDF' if cs else 'Outside-PDF anonymization', 'Soukromý originál se ověřuje pouze otiskem; veřejná kopie neobsahuje odstraněné údaje třetích osob a není označena za úřední kopii.' if cs else 'The private original is checked only by fingerprint; the public derivative contains none of the removed third-party data and is not labelled an official copy.'),
            ('Odesílání' if cs else 'Delivery', 'Systém nic automaticky neodeslal. Autor po své kontrole odeslal tři konečné dokumenty a schválil jejich zveřejnění v odeslané podobě; stav „doručeno“ se bez doručenky nepřiděluje.' if cs else 'The system filed nothing automatically. After human review, the author sent three final documents and approved publication in their sent form; no record is marked “delivered” without a delivery record.'),
        ],
        [5.0, 11.6]
    )

    doc.add_heading('10. Pokyny pro porotu' if cs else '10. Judge instructions', level=1)
    add_numbered(doc, [
        'Otevřít živý prototyp a nejprve zobrazit výchozí mapu bez nové listiny.' if cs else 'Open the live prototype and first display the baseline map without the new document.',
        'Spustit interní test nové listiny a zkontrolovat nový uzel OSZ Frýdek-Místek, časovou osu, lhůtu, semafor a kritická varování.' if cs else 'Run the internal new-document test and inspect the new Frýdek-Místek node, chronology, deadline result, traffic light, and critical warnings.',
        'Ve druhém testu načíst veřejnou anonymizovanou kopii PDF, porovnat stav paměti před a po vložení, zkontrolovat vytvořené řešení a semafor relevance a ověřit, že neznámé PDF připravenou analýzu nedostane.' if cs else 'In the second test, load the public anonymized PDF copy, compare memory before and after insertion, inspect the generated solution and relevance traffic light, and confirm that an unknown PDF receives no prepared analysis.',
        'Ověřit, že červená barva neznamená vinu a že systém vždy odděluje, co listina dokládá, nedokládá a co musí ověřit člověk.' if cs else 'Confirm that red does not mean guilt and that the system separates what the document proves, does not prove, and requires human verification.'
    ])
    add_links(doc, cs)

    doc.add_heading('11. Závěr testu' if cs else '11. Test conclusion', level=1)
    doc.add_paragraph(
        'Prototyp v tomto testu nevyrobil další souhrn. Provedl řízenou změnu institucionální paměti: rozpoznal nový uzel, zachoval důkazní meze, odmítl falešnou lhůtu, navrhl dokumentově odlišné doplnění, rozlišil relevantní a nerelevantní adresáty a zachytil drobné formulace i riziko chybného sloučení identity v čase. Přesně tam má AI Advocate přidanou hodnotu: ne v přesvědčivém tónu, ale v kontrolovatelnosti změny.'
        if cs else
        'The prototype did not merely produce another summary. It performed a controlled update of institutional memory: it recognised a new node, preserved evidential limits, rejected a false deadline, proposed document-specific supplements, separated relevant from irrelevant recipients, and detected small wording errors and the risk of merging identity across time. That is the added value of AI Advocate: not persuasive tone, but a change that can be audited.'
    )
    add_callout(
        doc,
        'Čtyři kontrolní otázky' if cs else 'Four control questions',
        ('Co se stalo?  ·  Co to dokládá?  ·  Co stále nevíme?  ·  Co musí ověřit člověk?'
         if cs else
         'What happened?  ·  What does it prove?  ·  What remains unknown?  ·  What must a human verify?'),
        'blue'
    )


def build_language(language, output):
    doc = Document()
    configure_document(doc)
    core = doc.core_properties
    core.title = (
        'AI Advocate for the Poor — Protokol testování před podáním přihlášky'
        if language == 'cs'
        else 'AI Advocate for the Poor — Pre-Submission Test Protocol'
    )
    core.subject = 'OpenAI Build Week / Devpost submission 1098095'
    core.author = 'Mgr. Dušan Dvořák; AI Advocate for the Poor'
    core.keywords = 'AI Advocate for the Poor, Devpost, OpenAI Build Week, test protocol'
    add_protocol(doc, language)
    doc.save(output)
    print(output)


def build():
    OUT.mkdir(parents=True, exist_ok=True)
    build_language('cs', OUTPUT_CS)
    build_language('en', OUTPUT_EN)


if __name__ == '__main__':
    build()
