"""Build the reviewed public v3.2 reopening outputs.

The source PDF filed by Dušan Dvořák is copied byte-for-byte by the release
workflow.  This builder creates only a separate addendum and three bounded
working documents.  It deliberately keeps official findings, party statements,
system comparisons, missing evidence, and proposed next steps separate.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Mm, Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
OUT = REPO / "work" / "generated-docs"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
BLACK = RGBColor(20, 24, 28)
MUTED = RGBColor(88, 96, 105)
LIGHT_BLUE = "EAF2F8"
LIGHT_RED = "FDECEC"
LIGHT_AMBER = "FFF4D6"
LIGHT_GREEN = "EAF6EC"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


@dataclass(frozen=True)
class SourceRow:
    date: str
    source: str
    reference: str
    quotation: str
    meaning: str
    limit: str


SHARED_THC_EVIDENCE: tuple[SourceRow, ...] = (
    SourceRow(
        "25. 6. 2026",
        "Kriminalistický ústav Policie ČR",
        "KU-4139-5/ČJ-2026-2305KM",
        "„Jednotlivé OKTE KŘP si standardní operační postupy (dále jen ‚SOP‘) tvoří samy.“",
        "Přímý úřední podklad k tomu, že konkrétní SOP vznikají na jednotlivých OKTE.",
        "Táž listina zároveň uvádí, že základní pravidla jsou sjednocená a že odběr byl sjednocen již před rokem 2021. Sama proto neprokazuje chybu konkrétního měření.",
    ),
    SourceRow(
        "25. 6. 2026",
        "Kriminalistický ústav Policie ČR",
        "KU-4139-5/ČJ-2026-2305KM",
        "„Způsob uvádění nejistot měření je ponecháno čistě na každé laboratoři a na jejich nastaveném systému kvality.“",
        "Přímý úřední podklad k tomu, že práci s nejistotou je nutné ověřit u konkrétní laboratoře a konkrétního posudku.",
        "Listina neříká, že nejistota nebyla v jednotlivém případě validována nebo zohledněna.",
    ),
    SourceRow(
        "27. 2. 2025",
        "Krajské státní zastupitelství v Ostravě",
        "2 KZT 59/2025-62",
        "„státní zástupkyně uvedla, že nedisponuje předpisy, kterými se tato pracoviště řídí při odborném zkoumání rostlin konopí“",
        "Přímá vazba mezi konkrétní trestní věcí G. F./J. K. a chybějícím předpisem ve spise státní zástupkyně.",
        "Nedisponování předpisem není důkazem jeho neexistence ani nesprávnosti výsledku 3,7 % THC.",
    ),
    SourceRow(
        "29. 7. 2025",
        "Vrchní soud v Praze",
        "11 To 88/2024-2990",
        "„z popisu skutku uvedeného ve výroku o vině nelze rozklíčovat to, zda se tedy jednalo o konopí tzv. zakázané či technické“",
        "Nezávislý soudní srovnávací pramen ukazuje, že druh konopí a význam chemických stop musí být z rozsudku přezkoumatelný.",
        "Jde o jinou trestní věc. Usnesení neřeší metodiku ve zdejších kauzách a nenařizuje zproštění.",
    ),
    SourceRow(
        "29. 7. 2025",
        "Vrchní soud v Praze",
        "11 To 88/2024-2990",
        "„Soud by tedy neměl jen ‚přebírat‘ výsledky do výroku rozsudku, ale měl by i znalecký posudek podrobit určité vlastní kontrole a vyhodnocení“",
        "Srovnávací soudní požadavek na skutečnou kontrolu znaleckého podkladu.",
        "Pozdější soudní názor není sám o sobě novou skutkovou okolností ve smyslu § 278 odst. 1 trestního řádu.",
    ),
)


INSTITUTIONAL_CHRONOLOGY: tuple[SourceRow, ...] = (
    SourceRow(
        "17. 11. 2024",
        "stížnost adresovaná ministrovi spravedlnosti a NSZ",
        "MSP-674/2024-ODKA-SPZ; 1 NZZ 105/2024",
        "„Stěžovatel proto žádá, aby nejvyšší státní zástupce společně (ev. samostatně) a ministr spravedlnosti po prozkoumání této stížnosti rozhodli…“",
        "Dokládá obsah a označení podání, v němž byla soustředěna metodická námitka THC.",
        "Jde o podání účastníka; bez samostatné doručenky neprokazuje přesný okamžik doručení ani správnost všech tvrzení.",
    ),
    SourceRow(
        "11. 3. 2025",
        "podání prezidentu republiky a předsedovi vlády",
        "bez č. j. odesílatele",
        "„Podávám tímto veřejně čestné prohlášení … o zmocnění k podání návrhů na zastavení trestních řízení (abolicí), milostí a rehabilitací…“",
        "Dokládá obsah a adresáty návrhů, včetně větve G. F./J. K., manželů K. a Dušana Dvořáka.",
        "Jde o podání účastníka; samo neprokazuje rozhodnutí prezidenta ani premiéra.",
    ),
    SourceRow(
        "14. 3. 2025",
        "dodatek prezidentu republiky Petru Pavlovi a předsedovi vlády Petru Fialovi",
        "bez č. j. odesílatele",
        "„Dodatek k návrhům ze dne 11. března 2025 na udělení abolice, milostí a rehabilitací“",
        "Rozšiřuje doloženou chronologii informování obou ústavních činitelů.",
        "Jde o tvrzení a přílohy podatele, nikoli o přijetí jejich skutkového nebo právního hodnocení adresáty.",
    ),
    SourceRow(
        "19. 3. 2025",
        "Úřad vlády České republiky",
        "11473/2025-UVCR",
        "„panu předsedovi vlády ČR byla doručena Vaše žádost o udělení milosti formou abolice pro členy aliance …“",
        "Přímá úřední odpověď potvrzuje doručení žádosti předsedovi vlády a vymezuje jeho roli u abolice; jména soukromých osob jsou ve veřejném výstupu vypuštěna.",
        "Nevyjadřuje se k důvodnosti trestních námitek ani k výsledku žádosti.",
    ),
    SourceRow(
        "6. 3. 2026",
        "datová schránka Kanceláře prezidenta republiky",
        "ID zprávy 1660532231",
        "„Stav zprávy: Dodaná … Datum a čas: 6. 3. 2026 v 14:06:51“",
        "Přímá doručenka potvrzuje dodání balíku KPR; seznam příloh zahrnuje dokumenty k THC, KÚ a NCOZ.",
        "Doručenka dokládá dodání, nikoli souhlas prezidenta s obsahem nebo jeho věcné posouzení.",
    ),
    SourceRow(
        "17. 6. 2026",
        "doplnění prezidentu republiky",
        "KPR 4093/2026",
        "„DOPLNĚNÍ ŽÁDOSTI O MILOST, AMNESTII, ABOLICI A REHABILITACI ve věci Mgr. Dušana Dvořáka a osob odsouzených za konopí“",
        "Dokládá obsah dalšího podání a jeho přiřazení ke spisu KPR.",
        "Samotný stejnopis podání neprokazuje výsledek ani úplnost spisu KPR.",
    ),
    SourceRow(
        "14. 7. 2026",
        "předžalobní výzva Nejvyššímu státnímu zastupitelství",
        "v textu související sp. zn. 6 NZN 1737/2026",
        "„PŘEDŽALOBNÍ VÝZVA k odstranění trvajícího stavu nepřezkoumatelnosti a roztříštěného vyřizování podání“",
        "Souhrnné podání spojuje data, č. j./sp. zn. a metodickou chronologii pro další případové návrhy.",
        "Jde o tvrzení podatele; listina sama nepotvrzuje správnost celé chronologie ani pochybení NSZ.",
    ),
    SourceRow(
        "15. 7. 2026",
        "stížnost ministrovi spravedlnosti JUDr. Jeronýmu Tejcovi",
        "MSP-162/2026-ODKA-SPZ; MSP-19/2026-ODKA-ROZ",
        "„Podávám k rukám ministra spravedlnosti stížnost na nečinnost Ministerstva spravedlnosti při věcném vyřízení opakovaných… podnětů“",
        "Dokládá přesný obsah, adresáta a dvě uvedené ministerské spisové větve.",
        "Jde o podání účastníka; neprokazuje nečinnost ani její protiprávnost bez úplných ministerských spisů.",
    ),
    SourceRow(
        "15. a 17. 7. 2026 / odpověď 21. 7. 2026",
        "podání ministrovi vnitra Mgr. Lubomíru Metnarovi a odpověď Ministerstva vnitra",
        "MV-108890-3/TP-2026; MV-97289/TP-2026; MV-114818-2/TP-2026",
        "„Ministerstvo vnitra dne 15. a 17. 7. 2026 obdrželo Vaše podání“",
        "Přímá úřední odpověď potvrzuje přijetí dvou červencových podání ministerstvem.",
        "Odpověď uvádí nepříslušnost; neposuzuje správnost laboratorních výsledků ani tvrzenou nezákonnost.",
    ),
)


LEGAL_SOURCES = (
    (
        "§§ 277–278 trestního řádu",
        "https://e-sbirka.gov.cz/sb/1961/141",
        "Obnova vyžaduje nové skutečnosti nebo důkazy dříve soudu neznámé, které by mohly vést k jinému rozhodnutí; procesní použitelnost musí být ověřena podle úplného spisu.",
    ),
    (
        "Tpjn 302/2012, Nejvyšší soud",
        "https://sbirka.nsoud.cz/sbirka/6773/",
        "Později zjištěný právní názor sám nestačí; navrhovaná nova se hodnotí ve vztahu k původním skutkovým zjištěním.",
    ),
    (
        "8 Tz 25/2015, Nejvyšší soud",
        "https://sbirka.nsoud.cz/sbirka/7413/",
        "Soud porovnává dokazování v původním řízení s novými důkazy předloženými v řízení o obnově.",
    ),
    (
        "IV. ÚS 1140/18, Ústavní soud",
        "https://nalus.usoud.cz/Search/GetText.aspx?sz=4-1140-18_1",
        "Veřejný prvotní pramen pro přesnou procesní cestu M. K./J. K.; nenahrazuje jejich úplný trestní spis.",
    ),
)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Sequence[int], indent=120):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name, size, color, bold in (
        ("Source Citation", 9, MUTED, False),
        ("Status Label", 9, DARK_BLUE, True),
    ):
        if style_name not in styles:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(4)


def add_page_number(paragraph):
    run = paragraph.add_run("Strana ")
    set_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "586069")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    r_pr.extend([color, size])
    text = OxmlElement("w:t")
    text.text = "1"
    r.append(r_pr)
    r.append(text)
    field.append(r)
    paragraph._p.append(field)


def configure_document(doc: Document, short_title: str, privacy: str):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)
    section.header_distance = Mm(9)
    section.footer_distance = Mm(9)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"AI Advocate for the Poor · v3.2 · {short_title}")
    set_font(run, size=8.5, bold=True, color=MUTED)
    run2 = p.add_run(f"\n{privacy}")
    set_font(run2, size=8, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(fp)

    props = doc.core_properties
    props.title = short_title
    props.subject = "AI Advocate for the Poor v3.2 — source-bounded legal document output"
    props.author = "Mgr. Dušan Dvořák; prepared with Codex assistance"
    props.keywords = "obnova řízení, THC, důkazní paměť, pracovní výstup"
    props.comments = "Generated from reviewed sources; requires human legal review before filing."


def add_bottom_rule(paragraph, color="2E74B5", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_title_block(doc: Document, kicker: str, title: str, subtitle: str, metadata: Sequence[tuple[str, str]]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(kicker.upper())
    set_font(r, size=9, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, size=21, bold=True, color=BLACK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(subtitle)
    set_font(r, size=13, color=MUTED)

    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        lr = p.add_run(f"{label}: ")
        set_font(lr, bold=True)
        vr = p.add_run(value)
        set_font(vr)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(10)
    add_bottom_rule(rule)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    relationship_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(r_pr)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_callout(doc: Document, label: str, text: str, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_table_geometry(table, [9360])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    lr = p.add_run(f"{label.upper()}\n")
    set_font(lr, size=9, bold=True, color=DARK_BLUE)
    tr = p.add_run(text)
    set_font(tr, size=11, bold=True, color=BLACK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)
    return table


def ensure_numbering(doc: Document):
    numbering = doc.part.numbering_part.element
    existing = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    next_abstract = max(existing, default=-1) + 1
    existing_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    next_num = max(existing_num, default=0) + 1

    def make(fmt: str, text: str, font: str | None = None):
        nonlocal next_abstract, next_num
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(next_abstract))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        p_pr.append(tabs)
        p_pr.append(ind)
        level.extend([start, num_fmt, lvl_text, suff, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_pr.append(r_fonts)
            level.append(r_pr)
        abstract.append(level)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(next_num))
        abstract_id = OxmlElement("w:abstractNumId")
        abstract_id.set(qn("w:val"), str(next_abstract))
        num.append(abstract_id)
        numbering.append(num)
        value = next_num
        next_abstract += 1
        next_num += 1
        return value

    return make("bullet", "•", "Symbol"), make("decimal", "%1.")


def add_list_item(doc: Document, text: str, num_id: int, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_source_table(doc: Document, rows: Iterable[SourceRow], title="Důkazní matice"):
    doc.add_heading(title, level=2)
    caption = doc.add_paragraph(style="Source Citation")
    caption.paragraph_format.keep_with_next = True
    caption.add_run("Každý řádek odděluje přesnou citaci, její význam a důkazní limit.")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ("Datum · zdroj", "Č. j./sp. zn. · přesná citace", "Co dokládá", "Co nedokládá")
    widths = (1450, 2850, 2550, 2510)
    for cell, label in zip(table.rows[0].cells, headers):
        shade_cell(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        r = p.add_run(label)
        set_font(r, size=9, bold=True, color=DARK_BLUE)
    repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for item in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        values = (
            f"{item.date}\n{item.source}",
            f"{item.reference}\n{item.quotation}",
            item.meaning,
            item.limit,
        )
        for cell, value in zip(cells, values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r, size=8.5)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_legal_sources(doc: Document, bullets: int):
    doc.add_heading("Právní kontrolní rámec", level=2)
    for label, url, meaning in LEGAL_SOURCES:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.3)
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run("• ")
        set_font(r, bold=True, color=BLUE)
        add_hyperlink(p, label, url)
        r2 = p.add_run(f" — {meaning}")
        set_font(r2)


def add_institutional_appendix(doc: Document):
    doc.add_page_break()
    add_source_table(doc, INSTITUTIONAL_CHRONOLOGY, "Příloha A — vybraná chronologie informování státu")
    add_callout(
        doc,
        "Jak tuto přílohu číst",
        "Stejnopis podání dokládá, co podatel uvedl a komu jej adresoval. Doručenka nebo odpověď dokládá přijetí. Ani jedno samo o sobě neprokazuje pravdivost všech tvrzení, pochybení adresáta ani výsledek řízení.",
        LIGHT_AMBER,
    )


def add_signature_block(doc: Document, identity="Mgr. Dušan Dvořák"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.keep_together = True
    r = p.add_run(identity)
    set_font(r, bold=True)
    r = p.add_run("\n— pracovní elektronický návrh; před podáním zkontroluje a podepíše oprávněná osoba —")
    set_font(r, size=9, italic=True, color=MUTED)


def build_dd_addendum():
    doc = Document()
    configure_styles(doc)
    bullets, numbers = ensure_numbering(doc)
    configure_document(doc, "Dodatek Dušana Dvořáka", "VEŘEJNÉ · plná identita autora se zveřejňuje s jeho výslovným souhlasem")
    add_title_block(
        doc,
        "Oddělený pracovní dodatek — není součástí podání z 12. 7. 2026",
        "DODATEK K NÁVRHU NA POVOLENÍ OBNOVY ŘÍZENÍ",
        "Nově doložené listiny po 12. červenci 2026 a jejich přesně omezený význam",
        (
            ("Adresát", "Okresní soud v Prostějově, Havlíčkova 16, 797 09 Prostějov"),
            ("Navrhovatel", "Mgr. Dušan Dvořák, nar. 12. ledna 1962 v Olomouci; Edukativní konopná klinika, 798 55 Ospělov 6; ID DS r8u3nhx"),
            ("Věci", "2 T 104/2010; 2 T 65/2011 / 2 Nt 1257/2013; 2 Nt 1151/2014; 3 Nt 1151/2014"),
            ("Datum výstupu", "22. července 2026"),
            ("Status", "PRACOVNÍ DODATEK — dosud nepodán; před podáním nutná kontrola úplného spisu"),
        ),
    )
    add_callout(
        doc,
        "Hlavní závěr systému",
        "Po podání návrhu dne 12. 7. 2026 přibyly tři nové procesní uzly a jeden nově získaný srovnávací soudní pramen. Nejsilnější další krok není tvrdit, že tyto pozdější listiny samy dokazují důvod obnovy, ale přesně je spojit s historickým způsobem vzorkování, rozlišením legálního a zakázaného konopí a kontrolou konkrétních znaleckých výstupů v každém ze čtyř spisů.",
        LIGHT_RED,
    )

    doc.add_heading("I. Vztah k již podanému návrhu", level=1)
    doc.add_paragraph(
        "Navrhovatel dne 12. července 2026 podal společný návrh na povolení obnovy označených řízení. Tento dokument původní návrh nemění, nenahrazuje a zpětně nepřepisuje. Je samostatným dodatkem omezeným na listiny získané nebo doručené po jeho podání. Přesná původní listina zůstává zveřejněna beze změny a její digitální otisk SHA-256 je e59f9ffde1001d232bd27d1017911d31ba85113413a8aec768d31d2c001ce691."
    )
    doc.add_paragraph(
        "Navrhovatel zachovává svůj původní petit. Žádá, aby soud nové podklady četl spolu s původním návrhem, avšak u každého zvlášť posoudil, zda jde o důkaz historické skutečnosti existující již v době původního řízení, o později vzniklou procesní událost, nebo jen o srovnávací právní pramen."
    )

    doc.add_heading("II. Nové listiny po 12. červenci 2026", level=1)
    rows = (
        SourceRow(
            "14. 7. 2026",
            "předžalobní výzva NSZ",
            "související sp. zn. 6 NZN 1737/2026",
            "„PŘEDŽALOBNÍ VÝZVA k odstranění trvajícího stavu nepřezkoumatelnosti a roztříštěného vyřizování podání“",
            "Dokládá obsah, adresáta a souhrnnou chronologii tvrzení navrhovatele po podání obnovy.",
            "Je to podání navrhovatele, nikoli věcné rozhodnutí NSZ ani potvrzení všech tvrzení.",
        ),
        SourceRow(
            "20. 7. 2026",
            "Policie ČR, Územní odbor Prostějov",
            "KRPM-100092-2/ČJ-2026-1412UO",
            "„podání adresované Nejvyššímu státnímu zastupitelství v Brně bylo bez přijetí dalšího opatření uloženo“",
            "Přímo dokládá přijetí stejnopisu výzvy policií dne 14. 7. a jeho uložení bez dalšího opatření.",
            "Nejde o rozhodnutí NSZ; listina nevysvětluje, jak se podání k policii dostalo, ani věcně nehodnotí THC či obnovu.",
        ),
        SourceRow(
            "21. 7. 2026",
            "Krajské státní zastupitelství v Brně",
            "1 KZT 475/2026-32",
            "„Po opatření příslušných spisových materiálů bude věc přezkoumána a o výsledku budete vyrozuměn.“",
            "Otevírá přezkum postupu OSZ Prostějov ve věci 1 ZT 11/2010.",
            "Není výsledkem přezkumu. Záhlaví uvádí podání 10. 7. 2026, text 10. 7. 2025; datum musí orgán vyjasnit.",
        ),
        SourceRow(
            "21. 7. 2026",
            "Ministerstvo vnitra",
            "MV-114818-2/TP-2026",
            "„Ministerstvo vnitra dne 15. a 17. 7. 2026 obdrželo Vaše podání … Vámi uváděné skutečnosti se nevztahují k působnosti ministerstva“",
            "Potvrzuje přijetí dvou červencových podání a stanovisko ministerstva k vlastní působnosti.",
            "Nejde o odborné přezkoumání laboratorních výsledků ani o soudní závěr k nezákonnému zásahu.",
        ),
        SourceRow(
            "29. 7. 2025; nově získáno po 12. 7. 2026",
            "Vrchní soud v Praze",
            "11 To 88/2024-2990",
            "„Stěžejním důvodem pro zrušení napadeného rozsudku je však jeho naprostá nepřezkoumatelnost“",
            "Srovnávací soudní pramen z jiné živé věci; konkrétně vytýká laické chemické úvahy, nejasné rozlišení technického a zakázaného konopí a nedostatečnou kontrolu znalce.",
            "Usnesení je z roku 2025, není rozhodnutím ve věci navrhovatele a samo není novou historickou skutkovou okolností jeho původních řízení.",
        ),
    )
    add_source_table(doc, rows, "Důkazní status nových listin")

    doc.add_heading("III. Případové propojení nového soudního pramene", level=1)
    doc.add_paragraph(
        "Vrchní soud v Praze v usnesení 11 To 88/2024-2990 nezformuloval obecné pravidlo pro všechny konopné věci. Jeho konkrétní výtky však určují otázky, které lze přezkoumat v původních spisech navrhovatele: zda skutková věta rozlišila legální a zakázaný materiál; jak byl vybrán a homogenizován vzorek; co přesně znamenal naměřený údaj; zda bylo rozlišeno THC a THCA; jak byla zohledněna nejistota; a zda soud znalecký výstup skutečně kontroloval."
    )
    for item in (
        "Sp. zn. 2 T 104/2010 — vyžádat původní protokol odběru, úplný znalecký výstup, tehdy použitý SOP a přesnou pasáž rozsudku, která z výsledku THC dovozovala trestní význam.",
        "Sp. zn. 2 T 65/2011 / 2 Nt 1257/2013 — oddělit materiál a měření vztahující se k úrodě 2010 a porovnat je se skutečně použitým laboratorním postupem.",
        "Sp. zn. 2 Nt 1151/2014 — porovnat tvrzení převzaté v 1 NZO 84/2015-14 a 11 Tdo 181/2015 s konkrétním SOP, který měl být závazně přejímán.",
        "Sp. zn. 3 Nt 1151/2014 — doložit vlastní chemický důkaz, jeho vzorek, měřicí řetězec a přesný význam pro výrok dané věci; nepřenášet závěr z jiného spisu.",
    ):
        add_list_item(doc, item, bullets)

    doc.add_heading("IV. Návrh na doplnění potřebného šetření", level=1)
    doc.add_paragraph("Navrhovatel navrhuje, aby soud k již označeným důkazům doplnil zejména:")
    for item in (
        "ověřenou kopii usnesení Vrchního soudu v Praze ze dne 29. 7. 2025, č. j. 11 To 88/2024-2990, pouze jako srovnávací soudní pramen a vodítko pro rozsah kontroly znaleckého dokazování;",
        "v každém ze čtyř spisů přesný znalecký nebo odborný výstup, laboratorní zadání, protokol vzorkování, vážení a homogenizace, chromatogramy, výpočet výsledku, nejistotu měření a tehdy použitý SOP;",
        "od Kriminalistického ústavu a příslušného OKTE označení dokumentu, jímž byl v rozhodné době upraven právě použitý postup, a vysvětlení vztahu tohoto postupu k pozdějšímu rozkazu č. 54/2021;",
        "vyrozumění o výsledku přezkumu KSZ Brno pod sp. zn. 1 KZT 475/2026 a úplný spis, z něhož bude patrné, co bylo skutečně přezkoumáno;",
        "přesné procesní vysvětlení, proč policie obdržela podání adresované NSZ, a zda policie disponuje jiným souvisejícím podáním než stejnopisem předžalobní výzvy.",
    ):
        add_list_item(doc, item, numbers)

    add_source_table(doc, SHARED_THC_EVIDENCE, "V. Společné důkazní kandidáty a jejich hranice")
    add_legal_sources(doc, bullets)
    add_institutional_appendix(doc)

    doc.add_heading("Návrh procesního postupu", level=1)
    doc.add_paragraph(
        "Navrhovatel žádá, aby byl tento dodatek po podání založen ke všem čtyřem označeným věcem a aby soud před rozhodnutím o obnově provedl výše navržené šetření. Konečnou formulaci návrhu na rozhodnutí navrhovatel ponechá v souladu s již podaným petitem a s výsledkem doplněného dokazování."
    )
    add_callout(
        doc,
        "Procesní hranice",
        "Pozdější úřední reakce a jiný soudní rozsudek samy obvykle nenahrazují novou skutkovou okolnost existující již v době původního rozhodnutí. Těžiště dodatku proto spočívá ve vyžádání historických laboratorních podkladů, jejichž obsah mohl být původnímu soudu neznámý, a v případovém vysvětlení jejich možného vlivu.",
        LIGHT_AMBER,
    )
    add_signature_block(doc)
    path = OUT / "dusan-dvorak-addendum-draft-2026-07-22.docx"
    doc.save(path)
    return path


def build_lch_motion():
    doc = Document()
    configure_styles(doc)
    bullets, numbers = ensure_numbering(doc)
    configure_document(doc, "L. CH. — pracovní obnova", "VEŘEJNÁ ANONYMIZOVANÁ VERZE · soukromý návrh z roku 2022 se nezveřejňuje")
    add_title_block(
        doc,
        "Pracovní anonymizovaný návrh — před podáním doplnit originály",
        "L. CH. — NÁVRH NA POVOLENÍ OBNOVY ŘÍZENÍ",
        "Porovnání soukromého návrhu z roku 2022 se zdrojovanou pamětí k 22. červenci 2026",
        (
            ("Soud", "Krajský soud v Brně"),
            ("Napadená rozhodnutí", "KS Brno 28. 2. 2019, č. j. 50 T 7/2018-603; VS Olomouc 6. 11. 2019, č. j. 5 To 39/2019-707"),
            ("Veřejné označení", "L. CH. z dokumentu České televize Smoke (2025)"),
            ("Datum výstupu", "22. července 2026"),
            ("Status", "PRACOVNÍ NÁVRH — nikoli právní rada ani tvrzení, že jsou podmínky obnovy splněny"),
        ),
    )
    add_callout(
        doc,
        "Relevance 9/9 — extrém, hoří",
        "Návrh z roku 2022 už přesně napadal vztah 17,5 kg materiálu, 237,4 g označeného reprezentativního vzorku a 6,4 g použitých při expertize. Pozdější úřední podklady nyní umožňují žádat konkrétní tehdejší SOP, záznam vzorkování, homogenizaci, chromatogramy, nejistotu a rozlišení THC/THCA. Relevance je vysoká; úspěch se nepředpovídá.",
        LIGHT_RED,
    )
    doc.add_heading("I. Procesní identita a veřejná anonymizace", level=1)
    doc.add_paragraph(
        "Podklad vychází ze soukromého desetistránkového návrhu z roku 2022. Veřejná verze nepřebírá jméno, datum narození, adresu, věznici, podpis ani jiné osobní identifikátory. Zachovává však soud, datum a číslo jednací, protože bez nich nelze výstup ověřit."
    )
    p = doc.add_paragraph()
    p.add_run("Veřejný kontext: ").bold = True
    add_hyperlink(p, "Česká televize — Smoke (2025)", "https://www.ceskatelevize.cz/porady/16298026696-smoke/")
    p.add_run(". Toto mediální označení není důkazem důvodu obnovy.")

    doc.add_heading("II. Skutečnosti citované v soukromém návrhu z roku 2022", level=1)
    source_rows = (
        SourceRow(
            "28. 2. 2019",
            "Krajský soud v Brně — údaj citovaný v soukromém návrhu",
            "50 T 7/2018-603",
            "„Dle rozsudků totiž bylo na 1 rostlině konopí téměř 600 gramů překatrované sušiny“",
            "Vymezuje konkrétní spor o výnos a osobní přičitatelnost devíti rostlin.",
            "Bez originálu rozsudku nelze ověřit úplnou skutkovou větu, všechny důkazy ani kontext údaje.",
        ),
        SourceRow(
            "5. 12. 2016",
            "odborné vyjádření policie — citace v soukromém návrhu",
            "KRPB-173659-52/TČ-2016-061371-TI; KRPB-3821-1/KT-2016",
            "„Obálka obsahující reprezentativní vzorek … o hmotnosti 237,4 gramů … celková hmotnost … 17.500 gramů“",
            "Identifikuje původní laboratorní řetězec, který musí být znovu opatřen v úplném znění.",
            "Citace sama neprokazuje, kdo vzorek vybral, jak byl homogenizován ani zda byl postup odborně vadný.",
        ),
        SourceRow(
            "2022",
            "soukromý návrh L. CH.",
            "k 50 T 7/2018",
            "„z těchto … 237,4 gramů k expertíze bylo při expertíze použito 6,4 gramů“",
            "Přímo formuluje případovou kontrolní otázku k reprezentativnosti a přepočtu na 17,5 kg.",
            "Před podáním nové verze je nutné citaci ověřit proti celému odbornému vyjádření a protokolu o vzorkování.",
        ),
        SourceRow(
            "6. 11. 2019",
            "Vrchní soud v Olomouci — údaj citovaný v soukromém návrhu",
            "5 To 39/2019-707",
            "„odsouzen k trestu 8 let ve vězení“",
            "Identifikuje pravomocné rozhodnutí, proti němuž návrh z roku 2022 směřoval.",
            "Veřejný výstup nemá originál rozsudku; přesný výrok, právní moc a současný stav musí být doloženy.",
        ),
    )
    add_source_table(doc, source_rows)

    doc.add_heading("III. Nové důkazní kandidáty, které návrh z roku 2022 nemohl obsahovat", level=1)
    add_source_table(doc, SHARED_THC_EVIDENCE)
    doc.add_paragraph(
        "Rozhodující není obecná podobnost konopných věcí. Pro L. CH. je možný průnik jen tam, kde nový dokument pomůže zpětně ověřit právě odborné vyjádření KRPB-3821-1/KT-2016, výběr 237,4 g, použití 6,4 g, přepočet na 17,5 kg, zjištěných 3,65 % THC a osobní přičitatelnost devíti rostlin."
    )

    doc.add_heading("IV. Pracovní důvody návrhu", level=1)
    for item in (
        "Nově opatřený historický laboratorní podklad: vyžádat SOP účinný v prosinci 2016 a doklady, podle nichž byly vybrány, homogenizovány a analyzovány konkrétní položky. Teprve jejich obsah může být novým důkazem o dříve existující skutečnosti.",
        "Materiální význam: vysvětlit, zda by jiný závěr o reprezentativnosti, přepočtu hmotnosti, nejistotě nebo poměru THC/THCA mohl ovlivnit množství účinné látky, rozsah skutku, právní kvalifikaci nebo trest.",
        "Individuální rozsah: oddělit devět rostlin připisovaných L. CH. od dalších rostlin spoluobviněných a nepřenášet souhrnný výsledek bez doloženého vzorkovacího pravidla.",
        "Srovnávací rozhodnutí 11 To 88/2024-2990: použít pouze k formulaci kontrolních otázek, ne jako automatický důvod obnovy nebo důkaz nesprávnosti původního rozsudku.",
        "Právní námitky k povolení, čl. 39 Listiny a právu EU držet odděleně od skutečně nových skutkových důkazů; pozdější právní argument sám podmínku § 278 odst. 1 nenaplňuje.",
    ):
        add_list_item(doc, item, numbers)

    doc.add_heading("V. Důkazy, které je nutné opatřit před podáním", level=1)
    for item in (
        "ověřené úplné kopie rozsudků KS Brno 50 T 7/2018-603 a VS Olomouc 5 To 39/2019-707 včetně právní moci;",
        "úplné odborné vyjádření KRPB-3821-1/KT-2016, zadání, protokoly o vydání a zajištění, fotodokumentaci a řetězec nakládání se vzorky;",
        "SOP použitý dne 5. 12. 2016, validační dokumentaci, chromatogramy, kalibraci, výpočet 3,65 %, nejistotu a rozhodovací pravidlo;",
        "doklad, kdo označil 237,4 g za reprezentativní vzorek a podle jakého pravidla;",
        "aktuální procesní status L. CH. a jeho výslovné schválení konečného návrhu; údaj o podmíněném propuštění 14. 7. 2026 je nyní pouze autorem potvrzenou chronologií.",
    ):
        add_list_item(doc, item, bullets)

    add_legal_sources(doc, bullets)
    add_institutional_appendix(doc)
    doc.add_heading("VI. Pracovní petit", level=1)
    doc.add_paragraph(
        "Po doplnění uvedených originálů lze navrhnout, aby Krajský soud v Brně podle § 278 odst. 1 trestního řádu posoudil nové skutečnosti a důkazy ve spojení s původním spisem a podle výsledku rozhodl o obnově řízení, které skončilo rozhodnutími č. j. 50 T 7/2018-603 a 5 To 39/2019-707. Přesný rozsah zrušovaných výroků a konečné znění petitu musí schválit L. CH. a kvalifikovaný právník po kontrole celého spisu."
    )
    add_signature_block(doc, "L. CH. — identitu a podpis doplní oprávněná osoba")
    path = OUT / "lch-motion-to-reopen-working-draft-2026-07-22.docx"
    doc.save(path)
    return path


def build_mk_jk_motion():
    doc = Document()
    configure_styles(doc)
    bullets, numbers = ensure_numbering(doc)
    configure_document(doc, "M. K. / J. K. — pracovní obnova", "VEŘEJNÁ ANONYMIZOVANÁ VERZE · osobní a zdravotní údaje se nezveřejňují")
    add_title_block(
        doc,
        "Pracovní anonymizovaný návrh — vyžaduje souhlas dotčených osob",
        "M. K. / J. K. — NÁVRH NA POVOLENÍ OBNOVY ŘÍZENÍ",
        "Čtyři oddělené důkazní osy: chemie, výnos, digitální stopa a podmíněná cesta podle § 278 odst. 4",
        (
            ("Soud prvního stupně", "Krajský soud v Hradci Králové"),
            ("Procesní cesta", "KS Hradec Králové 27. 2. 2017, 9 T 5/2016-948; VS Praha 12. 6. 2017, 11 To 48/2017-1036; NS 13. 12. 2017, 11 Tdo 1499/2017-48; ÚS 9. 7. 2019, IV. ÚS 1140/18"),
            ("Datum výstupu", "22. července 2026"),
            ("Status", "PRACOVNÍ NÁVRH — žádný úspěch ani splnění § 278 odst. 4 se nepředjímá"),
        ),
    )
    add_callout(
        doc,
        "Relevance 9/9 — dvě osy je nutné držet odděleně",
        "První osa porovnává původní chemický a výnosový důkaz s pozdějšími úředními podklady a usnesením 11 To 88/2024-2990. Druhá, § 278 odst. 4, vznikne jen tehdy, prokáže-li pravomocný rozsudek, že soudce v původním řízení porušil povinnost jednáním zakládajícím trestný čin. Veřejně ověřená zpráva z 23. 1. 2026 popisuje odsouzení Ivana Elischera jako nepravomocné; podmínka proto zatím není doložena.",
        LIGHT_RED,
    )

    doc.add_heading("I. Ověřená procesní cesta", level=1)
    official_rows = (
        SourceRow(
            "12. 6. 2017",
            "Vrchní soud v Praze — shrnuto Ústavním soudem",
            "11 To 48/2017-1036",
            "„byl … v celém rozsahu zrušen rozsudek Krajského soudu … ze dne 27. 2. 2017 č. j. 9 T 5/2016-948. Nově bylo rozhodnuto…“",
            "Úřední pramen potvrzuje, že odsuzujícím pravomocným rozhodnutím byl nový rozsudek vrchního soudu.",
            "Veřejné usnesení ÚS nenahrazuje úplný rozsudek vrchního soudu ani jeho znalecké přílohy.",
        ),
        SourceRow(
            "13. 12. 2017",
            "Nejvyšší soud",
            "11 Tdo 1499/2017-48",
            "„tak, že obě dovolání odmítl“",
            "Potvrzuje výsledek dovolacího řízení.",
            "Citace neřeší, zda existují později zjištěná nova.",
        ),
        SourceRow(
            "9. 7. 2019",
            "Ústavní soud",
            "IV. ÚS 1140/18",
            "„Ústavní stížnost se odmítá.“",
            "Veřejný prvotní pramen umožňuje ověřit procesní cestu a skutkové shrnutí.",
            "Odmítnutí ústavní stížnosti nebrání návrhu na obnovu založenému na skutečně nových důkazech, ale samo jej ani nepodporuje.",
        ),
        SourceRow(
            "9. 7. 2019",
            "Ústavní soud — skutkové shrnutí",
            "IV. ÚS 1140/18, body 4 a 34",
            "„celkem 545 kusů rostlin … nejméně 28 855 gramů upotřebitelné sušiny“",
            "Vymezuje původní množstevní a výnosový závěr, proti němuž musí být nové důkazy konkrétně postaveny.",
            "Tento výrok není novým znaleckým posouzením Ústavního soudu; reprodukuje skutková zjištění obecných soudů.",
        ),
    )
    add_source_table(doc, official_rows)

    doc.add_heading("II. Osa A — chemie a měření THC", level=1)
    add_source_table(doc, SHARED_THC_EVIDENCE)
    doc.add_paragraph(
        "Aby byly společné podklady materiálně významné pro M. K./J. K., musí se propojit s konkrétním odborným vyjádřením zkušební laboratoře, vzorky zajištěnými při prohlídkách, tehdy použitým SOP a přesným způsobem, jak byly dílčí hodnoty přepočteny na množství čistého THC uvedené ve skutkové větě."
    )

    doc.add_heading("III. Osa B — výnos a množství upotřebitelné sušiny", level=1)
    doc.add_paragraph(
        "Ústavní soud v bodu 34 zaznamenal, že obecné soudy vycházely z odborného vyjádření Ing. Josefa Beneše a znaleckého posudku MUDr. Petra Petra, přičemž jedno vyjádření řešilo možný výnos po dorostení a druhé aktuálně zajištěné stopy. Pracovní návrh proto nesmí oba důkazní předměty směšovat. Novým důkazem by mohl být zejména nový nezávislý odborný posudek založený na původní dokumentaci, který konkrétně přezkoumá vstupní data, odrůdu, fázi růstu, podmínky pěstování a matematický přepočet."
    )

    doc.add_heading("IV. Osa C — digitální zdroje a individuální přičitatelnost", level=1)
    doc.add_paragraph(
        "Veřejné rozhodnutí IV. ÚS 1140/18 zachycuje námitky k fotografiím, datovým nosičům a zprávám. Pro obnovu nestačí námitky zopakovat. Je nutné získat nový forenzní podklad k původu souborů, metadatům, vlastnictví zařízení a vztahu konkrétní komunikace ke každé osobě. Usnesení 11 To 88/2024-2990 je zde pouze srovnávacím vodítkem, protože zdůrazňuje individuální hodnocení jednání každého obžalovaného."
    )

    doc.add_heading("V. Osa D — § 278 odst. 4 trestního řádu", level=1)
    doc.add_paragraph(
        "Ústavní soud v bodu 40 uvedl, že důvod obnovy podle § 278 odst. 4 by mohl vzniknout, byl-li by některý člen senátu pravomocně uznán vinným, že v původním řízení porušil své povinnosti jednáním zakládajícím trestný čin. Z předložených podkladů je doloženo, že členem senátu vrchního soudu byl JUDr. Ivan Elischer. Není však doloženo, že případné trestní jednání bylo pravomocně spojeno právě s rozhodováním ve věci 11 To 48/2017."
    )
    p = doc.add_paragraph(style="Source Citation")
    p.add_run("Aktuální veřejná kontrola: ")
    add_hyperlink(p, "ČT24, 23. 1. 2026 — opakovaný rozsudek označen jako nepravomocný", "https://ct24.ceskatelevize.cz/clanek/domaci/soud-uznal-i-napodruhe-vinnym-soudce-elischera-369569")
    p.add_run(". Před použitím § 278 odst. 4 je nutné opatřit pravomocný rozsudek, doložku právní moci a konkrétní skutkovou vazbu k původnímu řízení M. K./J. K.")

    doc.add_heading("VI. Důkazní návrhy před podáním", level=1)
    for item in (
        "úplný pravomocný rozsudek VS Praha 11 To 48/2017-1036 a rozsudek KS Hradec Králové 9 T 5/2016-948;",
        "původní chemické odborné vyjádření, všechny laboratorní položky, SOP, chromatogramy, nejistotu měření, odběr a homogenizaci;",
        "úplné odborné vyjádření Ing. Josefa Beneše a posudek MUDr. Petra Petra včetně vstupních dat a výpočtů;",
        "forenzní kopie a metadata digitálních nosičů, protokoly zajištění a přiřazení jednotlivých zpráv a fotografií;",
        "případný pravomocný rozsudek ve věci soudce Ivana Elischera a konkrétní důkaz jeho vztahu k rozhodování 11 To 48/2017;",
        "výslovný souhlas M. K. a J. K. s konečným zněním, jejich současný procesní stav a zmocnění osoby, která návrh podá.",
    ):
        add_list_item(doc, item, bullets)

    add_legal_sources(doc, bullets)
    add_institutional_appendix(doc)
    doc.add_heading("VII. Pracovní petit", level=1)
    doc.add_paragraph(
        "Po doplnění originálů lze navrhnout, aby Krajský soud v Hradci Králové podle § 278 odst. 1 trestního řádu porovnal označené nové skutečnosti a důkazy s původním spisem a rozhodl o povolení obnovy v rozsahu, který mohou skutečně ovlivnit. Důvod podle § 278 odst. 4 lze uplatnit pouze podmíněně po doložení pravomocného rozsudku a jeho konkrétní vazby k původnímu řízení."
    )
    add_signature_block(doc, "M. K. / J. K. — identity a podpisy doplní oprávněné osoby")
    path = OUT / "mk-jk-motion-to-reopen-working-draft-2026-07-22.docx"
    doc.save(path)
    return path


def build_gf_jk_memo():
    doc = Document()
    configure_styles(doc)
    bullets, numbers = ensure_numbering(doc)
    configure_document(doc, "G. F. / J. K. — procesní memorandum", "VEŘEJNÁ ANONYMIZOVANÁ VERZE · soukromé zdrojové soubory se nezveřejňují")
    add_title_block(
        doc,
        "Procesní memorandum — zatím nikoli návrh na obnovu",
        "G. F. / J. K. — CO JE NUTNÉ ZJISTIT PŘED VOLBOU OPRAVNÉHO PROSTŘEDKU",
        "Podmíněné zastavení, výsledek 3,7 % THC a chybějící konečný procesní stav",
        (
            ("Soud", "Okresní soud v Ostravě"),
            ("Uváděná věc", "usnesení ze dne 18. 6. 2025, č. j. 15 T 11/2025-122"),
            ("Souvisící policejní věc", "KRPT-202999/TČ-2024-070774; odborné vyjádření KRPT-2600-1/KT-2024"),
            ("Datum výstupu", "22. července 2026"),
            ("Status", "PRACOVNÍ MEMORANDUM — doložené podklady popisují podmíněné zastavení, nikoli odsouzení"),
        ),
    )
    add_callout(
        doc,
        "Relevance metodické osy 9/9; procesní cesta zatím neurčena",
        "Listina KSZ přímo spojuje výsledek 3,7 % THC s konkrétním odborným vyjádřením a současně zaznamenává, že státní zástupkyně neměla předpisy OKTE. To odůvodňuje získání laboratorního řetězce. Bez původního usnesení, právní moci a rozhodnutí po zkušební době však nelze bezpečně podat návrh na obnovu.",
        LIGHT_RED,
    )

    doc.add_heading("I. Co uvádějí úřední listiny", level=1)
    official_rows = (
        SourceRow(
            "27. 2. 2025",
            "Krajské státní zastupitelství v Ostravě",
            "2 KZT 59/2025-62",
            "„zkoumáno bylo 42,4 g rostlinné hmoty … 3,7 % THC … 1,6 g THC“",
            "Přímý úřední pramen k analyzovanému množství, výsledku a absolutní hmotnosti THC.",
            "Neobsahuje úplný SOP, chromatogramy, nejistotu ani oddělené výsledky všech zajištěných položek.",
        ),
        SourceRow(
            "27. 2. 2025",
            "Krajské státní zastupitelství v Ostravě",
            "2 KZT 59/2025-62",
            "„vzhledem k jejímu množství (celkem 8,2 g hmoty) nebylo stanoveno procento účinné látky ve vzorku“",
            "Dokládá, že u další položky nebylo procento THC kvantifikováno.",
            "Nevysvětluje samo, jak byl výsledek 3,7 % přiřazen jednotlivým rostlinám nebo celému skutku.",
        ),
        SourceRow(
            "27. 2. 2025",
            "Krajské státní zastupitelství v Ostravě",
            "2 KZT 59/2025-62",
            "„v dané věci nebyla zjištěna pochybení policejního orgánu vyžadující zásah dozorového státního zastupitelství“",
            "Přímý úřední závěr dohledové listiny.",
            "Nejde o soudní přezkum laboratorní metody a nevylučuje nové důkazy získané později.",
        ),
    )
    add_source_table(doc, official_rows)

    doc.add_heading("II. Co zatím pochází jen z pozdějších podání", level=1)
    party_rows = (
        SourceRow(
            "24. 6. 2025",
            "podání aliance",
            "odkaz na 15 T 11/2025-122",
            "„podle § 307, odst. 1 a 4 trestního řádu na dvanáct měsíců podmíněně pozastavil trestní řízení“",
            "Podání tvrdí podmíněné zastavení na dvanáct měsíců.",
            "Přesné znění, právní moc, podmínky a datum konce zkušební doby musí potvrdit originál usnesení.",
        ),
        SourceRow(
            "1. 3. 2026",
            "stížnost G. F./J. K.",
            "odkaz na sdělení soudu z 13. 1. 2026",
            "„šlo o pěstování celkem 6 kusů … a sušení dalšího 1 kusu … vše s obsahem 3,7 % THC“",
            "Podání zachycuje tvrzený pozdější popis 6 + 1 rostliny.",
            "Originál soudního sdělení není v kontrolovaném balíku; jiné listiny pracují s 3 + 1 nebo pěti rostlinami.",
        ),
        SourceRow(
            "24. 6. 2025",
            "podání aliance",
            "odkaz na jednání 18. 6. 2025",
            "„Soudkyně … potvrdila, že ověřila jako pravdivé tvrzení lékařů o účelnosti používání cannabisterapie“",
            "Podání uplatňuje léčebný účel.",
            "KSZ ve své listině uvádí opak; bez protokolu a usnesení nejde rozpor rozhodnout.",
        ),
    )
    add_source_table(doc, party_rows)

    doc.add_heading("III. Správné pořadí dalších kroků", level=1)
    steps = (
        "Opatřit usnesení OS Ostrava ze dne 18. 6. 2025, č. j. 15 T 11/2025-122, a doložku právní moci.",
        "Zjistit přesné datum začátku a konce zkušební doby a opatřit rozhodnutí vydané po jejím skončení podle § 308 trestního řádu.",
        "Sjednotit počet rostlin pomocí protokolu o vydání věci, odborného vyjádření, návrhu na potrestání a soudního usnesení; rozdíl 3 + 1 / 5 / 6 + 1 zatím není prokázanou chybou.",
        "Opatřit úplné odborné vyjádření KRPT-2600-1/KT-2024, použitý SOP, vzorkování, homogenizaci, chromatogramy, nejistotu a výsledky jednotlivých položek.",
        "Teprve podle konečného procesního stavu zvolit prostředek: žádost o záznam osvědčení, jiný procesní návrh, nebo podmíněně obnovu rozhodnutí uvedeného v § 277 trestního řádu.",
    )
    for item in steps:
        add_list_item(doc, item, numbers)

    add_source_table(doc, SHARED_THC_EVIDENCE, "IV. Sdílené důkazní kandidáty")
    add_legal_sources(doc, bullets)
    add_institutional_appendix(doc)
    doc.add_heading("V. Podmíněná osnova budoucího podání", level=1)
    doc.add_paragraph(
        "Jestliže originální listiny potvrdí pravomocné skončení způsobem, pro který je obnova přípustná, lze připravit návrh oddělující: (1) napadené rozhodnutí; (2) přesný původní výsledek 3,7 % THC a jeho důkazní řetězec; (3) skutečně nové historické důkazy k tehdy použitému SOP; (4) jejich možný vliv na rozhodnutí; a (5) přesný petit. Do té doby by označení dokumentu jako návrhu na obnovu vytvářelo nepravdivý dojem o procesním stavu."
    )
    add_signature_block(doc, "G. F. / J. K. — identity a podpisy doplní oprávněné osoby")
    path = OUT / "gf-jk-procedural-memorandum-2026-07-22.docx"
    doc.save(path)
    return path


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    paths = [build_dd_addendum(), build_lch_motion(), build_mk_jk_motion(), build_gf_jk_memo()]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
